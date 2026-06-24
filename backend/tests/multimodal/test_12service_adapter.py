"""P4-7-W1: 12 service shared multimodal adapter tests (smoke: 6 modalities / 3 outputs)."""
from __future__ import annotations

import base64
import io
import math
import os

import pytest

from common.multimodal_adapter import (
    AdapterRequest,
    MultimodalAdapter,
    MultimodalStore,
    OUTPUT_TEXT,
    OUTPUT_JSON,
    OUTPUT_MULTIMODAL,
    build_multimodal_router,
)
from imdf.multimodal.parser import (
    MODALITY_TEXT,
    MODALITY_IMAGE,
    MODALITY_AUDIO,
    MODALITY_VIDEO,
    MODALITY_DOCUMENT,
    MODALITY_EMAIL,
)


@pytest.fixture
def adapter() -> MultimodalAdapter:
    return MultimodalAdapter(service_id="search_service")


def test_text_modality(adapter: MultimodalAdapter):
    resp = adapter.process(AdapterRequest(
        service_id="search_service", modality=MODALITY_TEXT,
        text="hello world", output_kind=OUTPUT_TEXT,
    ))
    assert resp.text == "hello world"
    assert resp.modality == MODALITY_TEXT
    assert resp.embedding and len(resp.embedding) == 1024


def test_image_modality(adapter: MultimodalAdapter):
    from PIL import Image
    img = Image.new("RGB", (8, 8), (255, 0, 0))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    b64 = base64.b64encode(buf.getvalue()).decode("ascii")
    resp = adapter.process(AdapterRequest(
        service_id="search_service", modality=MODALITY_IMAGE,
        base64=b64, filename="red.png", mime_type="image/png",
        output_kind=OUTPUT_JSON,
    ))
    assert resp.modality == MODALITY_IMAGE
    assert resp.structured.get("modality") == "image"
    assert resp.structured.get("dimensions") == [8, 8]


def test_audio_modality(adapter: MultimodalAdapter):
    import struct
    import wave
    sr = 8000
    n = 800
    with io.BytesIO() as buf:
        with wave.open(buf, "wb") as w:
            w.setnchannels(1)
            w.setsampwidth(2)
            w.setframerate(sr)
            w.writeframes(b"\x00\x00" * n)
        data = buf.getvalue()
    b64 = base64.b64encode(data).decode("ascii")
    resp = adapter.process(AdapterRequest(
        service_id="search_service", modality=MODALITY_AUDIO,
        base64=b64, filename="x.wav", mime_type="audio/wav",
        output_kind=OUTPUT_JSON,
    ))
    assert resp.modality == MODALITY_AUDIO
    assert resp.structured.get("modality") == "audio"


def test_video_modality(adapter: MultimodalAdapter):
    import numpy as np  # type: ignore
    cv2 = pytest.importorskip("cv2")
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".avi", delete=False) as tmp:
        path = tmp.name
    try:
        fourcc = cv2.VideoWriter_fourcc(*"MJPG")
        writer = cv2.VideoWriter(path, fourcc, 10.0, (32, 32))
        for i in range(4):
            writer.write(np.zeros((32, 32, 3), dtype=np.uint8))
        writer.release()
        with open(path, "rb") as f:
            data = f.read()
    finally:
        try:
            os.unlink(path)
        except Exception:  # noqa: BLE001
            pass
    b64 = base64.b64encode(data).decode("ascii")
    resp = adapter.process(AdapterRequest(
        service_id="search_service", modality=MODALITY_VIDEO,
        base64=b64, filename="x.avi", mime_type="video/x-msvideo",
        output_kind=OUTPUT_MULTIMODAL,
    ))
    assert resp.modality == MODALITY_VIDEO
    assert resp.multimodal.get("handler", {}).get("modality") == "video"


def test_document_modality(adapter: MultimodalAdapter):
    from docx import Document
    d = Document()
    d.add_heading("Test", level=1)
    d.add_paragraph("Some content here for testing multimodal adapter.")
    buf = io.BytesIO()
    d.save(buf)
    b64 = base64.b64encode(buf.getvalue()).decode("ascii")
    resp = adapter.process(AdapterRequest(
        service_id="search_service", modality=MODALITY_DOCUMENT,
        base64=b64, filename="x.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        output_kind=OUTPUT_JSON,
    ))
    assert resp.modality == MODALITY_DOCUMENT
    assert resp.structured.get("modality") == "document"


def test_email_modality(adapter: MultimodalAdapter):
    eml = (
        b"From: a@b.com\r\nTo: c@d.com\r\n"
        b"Subject: P4-7-W1 test\r\nDate: now\r\n"
        b"Content-Type: text/plain; charset=utf-8\r\n\r\n"
        b"Hello there.\r\n"
    )
    b64 = base64.b64encode(eml).decode("ascii")
    resp = adapter.process(AdapterRequest(
        service_id="search_service", modality=MODALITY_EMAIL,
        base64=b64, filename="x.eml", mime_type="message/rfc822",
        output_kind=OUTPUT_JSON,
    ))
    assert resp.modality == MODALITY_EMAIL
    assert "P4-7-W1" in resp.structured.get("subject", "")


def test_routing_table(adapter: MultimodalAdapter):
    """Modality → service routing should match the spec."""
    assert adapter.route(MODALITY_IMAGE) == "user_service"
    assert adapter.route(MODALITY_VIDEO) == "asset_service"
    assert adapter.route(MODALITY_AUDIO) == "annotation_service"
    assert adapter.route(MODALITY_DOCUMENT) == "search_service"
    assert adapter.route(MODALITY_EMAIL) == "notification_service"


def test_build_router_endpoints():
    """The shared router exposes 4 endpoints."""
    from fastapi import FastAPI
    app = FastAPI()
    app.include_router(build_multimodal_router(
        service_id="user_service",
        adapter=MultimodalAdapter(service_id="user_service"),
    ))
    paths = sorted({r.path for r in app.routes if hasattr(r, "path")})
    assert "/api/v1/multimodal/health" in paths
    assert "/api/v1/multimodal/modalities" in paths
    assert "/api/v1/multimodal/process" in paths
    assert "/api/v1/multimodal/records" in paths
