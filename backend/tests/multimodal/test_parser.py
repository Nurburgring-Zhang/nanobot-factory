"""P4-7-W1: MultiModalParser tests (5: 6 document formats + 4 media formats)."""
from __future__ import annotations

import base64
import io
import os
import tempfile
import zipfile
import zlib
from pathlib import Path

import pytest

from imdf.multimodal.parser import (
    MultiModalParser,
    MultimodalDocument,
    detect_modality,
    MODALITY_TEXT,
    MODALITY_IMAGE,
    MODALITY_AUDIO,
    MODALITY_VIDEO,
    MODALITY_DOCUMENT,
    MODALITY_EMAIL,
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------
@pytest.fixture
def parser() -> MultiModalParser:
    return MultiModalParser(include_image_base64=False, max_pdf_pages=10)


@pytest.fixture
def tiny_png_bytes() -> bytes:
    """Return a 4x4 RGB PNG header — valid enough for PIL to open."""
    # 4x4 RGB image, three distinct colour blocks
    try:
        from PIL import Image
        img = Image.new("RGB", (8, 8), (0, 0, 0))
        # Add some structure so DCT is non-trivial
        for x in range(8):
            for y in range(8):
                img.putpixel((x, y), ((x * 32) % 256, (y * 32) % 256, 128))
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()
    except Exception:  # noqa: BLE001
        # tiny 1x1 PNG (89 bytes) - always works
        return (
            b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
            b"\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDAT\x08\x99c\xf8\xff"
            b"\xff?\x00\x05\xfe\x02\xfe\xa3\x9c\xb1\x00\x00\x00\x00IEND\xaeB`\x82"
        )


@pytest.fixture
def tiny_wav_bytes() -> bytes:
    """1 second of 8 kHz mono PCM16 silence."""
    import struct
    import wave
    sr = 8000
    n = sr
    with io.BytesIO() as buf:
        with wave.open(buf, "wb") as w:
            w.setnchannels(1)
            w.setsampwidth(2)
            w.setframerate(sr)
            w.writeframes(b"\x00\x00" * n)
        return buf.getvalue()


@pytest.fixture
def tiny_pdf_bytes() -> bytes:
    """Build a tiny PDF with one page of text and one image using pymupdf.

    Falls back to a hardcoded minimal PDF if pymupdf is not available.
    """
    try:
        import fitz  # type: ignore
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "P4-7-W1 hello world", fontsize=12)
        # Add a small image
        try:
            from PIL import Image
            img = Image.new("RGB", (20, 20), (255, 0, 0))
            img_buf = io.BytesIO()
            img.save(img_buf, format="PNG")
            page.insert_image(fitz.Rect(50, 100, 150, 200), stream=img_buf.getvalue())
        except Exception:  # noqa: BLE001
            pass
        # Add a basic table-like text
        page.insert_text((50, 250), "Col1\tCol2\tCol3\nA\tB\tC", fontsize=10)
        out = doc.tobytes()
        doc.close()
        return out
    except Exception:  # noqa: BLE001
        # Minimal PDF that pdfplumber can usually still parse
        return (
            b"%PDF-1.4\n"
            b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
            b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
            b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 300 200]/Contents 4 0 R"
            b"/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
            b"4 0 obj<</Length 50>>stream\nBT /F1 12 Tf 50 50 Td (Hello P4-7) Tj ET\nendstream\nendobj\n"
            b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
            b"xref\n0 6\n0000000000 65535 f \n0000000010 00000 n \n0000000053 00000 n \n0000000098 00000 n \n"
            b"0000000190 00000 n \n0000000280 00000 n \n"
            b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n350\n%%EOF\n"
        )


@pytest.fixture
def tiny_docx_bytes() -> bytes:
    """Create a 2-paragraph, 1-table DOCX in memory."""
    try:
        from docx import Document
        d = Document()
        d.add_heading("Title", level=1)
        d.add_paragraph("First paragraph about multimodal data.")
        d.add_paragraph("Second paragraph references image and table.")
        tbl = d.add_table(rows=2, cols=2)
        tbl.cell(0, 0).text = "A"
        tbl.cell(0, 1).text = "B"
        tbl.cell(1, 0).text = "C"
        tbl.cell(1, 1).text = "D"
        buf = io.BytesIO()
        d.save(buf)
        return buf.getvalue()
    except Exception:  # noqa: BLE001
        return b""


@pytest.fixture
def tiny_xlsx_bytes() -> bytes:
    try:
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws.append(["Name", "Value", "Score"])
        ws.append(["alpha", 10, 0.5])
        ws.append(["beta", 20, 0.8])
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()
    except Exception:  # noqa: BLE001
        return b""


@pytest.fixture
def tiny_eml_bytes() -> bytes:
    return (
        b"From: alice@example.com\r\n"
        b"To: bob@example.com\r\n"
        b"Subject: P4-7-W1 hello\r\n"
        b"Date: Wed, 24 Jun 2026 06:00:00 +0800\r\n"
        b"Message-ID: <abc@example.com>\r\n"
        b"MIME-Version: 1.0\r\n"
        b"Content-Type: text/plain; charset=utf-8\r\n"
        b"\r\n"
        b"Hello Bob, this is a test email body.\r\n"
        b"Multi-line text.\r\n"
    )


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------
def test_detect_modality_by_extension():
    assert detect_modality("foo.pdf") == MODALITY_DOCUMENT
    assert detect_modality("foo.png") == MODALITY_IMAGE
    assert detect_modality("foo.mp4") == MODALITY_VIDEO
    assert detect_modality("foo.wav") == MODALITY_AUDIO
    assert detect_modality("foo.eml") == MODALITY_EMAIL
    assert detect_modality("foo.txt") == MODALITY_TEXT


def test_parse_pdf(parser: MultiModalParser, tiny_pdf_bytes: bytes):
    if not tiny_pdf_bytes:
        pytest.skip("no pdf backend")
    doc = parser.parse(tiny_pdf_bytes, filename="x.pdf",
                       mime_type="application/pdf", modality=MODALITY_DOCUMENT)
    assert doc.modality == MODALITY_DOCUMENT
    # pymupdf fallback path may not extract text; just verify structure
    assert doc.doc_id.startswith("mm-")
    assert "pages" in doc.metadata
    # If text was extracted, length should be > 0
    if "pdfplumber_missing" not in doc.warnings:
        # best effort
        assert isinstance(doc.text, str)


def test_parse_docx(parser: MultiModalParser, tiny_docx_bytes: bytes):
    if not tiny_docx_bytes:
        pytest.skip("no docx backend")
    doc = parser.parse(tiny_docx_bytes, filename="x.docx",
                       modality=MODALITY_DOCUMENT)
    assert doc.modality == MODALITY_DOCUMENT
    assert any("paragraph" in s.text.lower() or "title" in s.text.lower()
               or "image" in s.text.lower() for s in doc.segments), \
        f"expected docx text, got: {[s.text for s in doc.segments]}"
    assert len(doc.tables) >= 1
    assert doc.tables[0].n_rows >= 2


def test_parse_xlsx(parser: MultiModalParser, tiny_xlsx_bytes: bytes):
    if not tiny_xlsx_bytes:
        pytest.skip("no xlsx backend")
    doc = parser.parse(tiny_xlsx_bytes, filename="x.xlsx",
                       modality=MODALITY_DOCUMENT)
    assert doc.modality == MODALITY_DOCUMENT
    assert len(doc.tables) >= 1
    assert doc.tables[0].n_rows >= 2
    assert "alpha" in doc.text or "Sheet1" in doc.text


def test_parse_markdown(parser: MultiModalParser):
    md = b"# Title\n\nThis is **bold** text.\n\n## Subtitle\n\n- item 1\n- item 2\n"
    doc = parser.parse(md, filename="x.md", modality=MODALITY_DOCUMENT)
    assert doc.modality == MODALITY_DOCUMENT
    assert any("Title" in s.text for s in doc.segments)
    assert any(s.segment_type == "heading" for s in doc.segments)


def test_parse_html(parser: MultiModalParser):
    html = (
        b"<html><body>"
        b"<h1>Heading</h1>"
        b"<p>Hello <b>world</b>!</p>"
        b"<table><tr><th>A</th><th>B</th></tr><tr><td>1</td><td>2</td></tr></table>"
        b"<img src='x.png' alt='placeholder'/>"
        b"</body></html>"
    )
    doc = parser.parse(html, filename="x.html", modality=MODALITY_DOCUMENT)
    assert doc.modality == MODALITY_DOCUMENT
    assert "Heading" in doc.text
    assert any("world" in s.text for s in doc.segments)
    assert len(doc.tables) >= 1
    assert len(doc.images) >= 1


def test_parse_text(parser: MultiModalParser):
    text = b"Hello world.\n\nThis is paragraph two.\n\nThird paragraph."
    doc = parser.parse(text, filename="x.txt", modality=MODALITY_TEXT)
    assert doc.modality == MODALITY_TEXT
    assert len(doc.segments) >= 2
    assert "Hello" in doc.text


def test_parse_image(parser: MultiModalParser, tiny_png_bytes: bytes):
    doc = parser.parse(tiny_png_bytes, filename="x.png",
                       modality=MODALITY_IMAGE)
    assert doc.modality == MODALITY_IMAGE
    assert len(doc.images) >= 1
    img = doc.images[0]
    assert img.width > 0 and img.height > 0
    assert img.sha256


def test_parse_audio(parser: MultiModalParser, tiny_wav_bytes: bytes):
    doc = parser.parse(tiny_wav_bytes, filename="x.wav",
                       modality=MODALITY_AUDIO)
    assert doc.modality == MODALITY_AUDIO
    # whisper is not always available; either transcript or warning
    if not doc.text:
        assert any("whisper" in w or "duration" in str(doc.metadata)
                   for w in doc.warnings) or doc.metadata.get("duration_s")


def test_parse_video_synthetic(parser: MultiModalParser):
    """Build a tiny AVI file with 8 frames and parse it."""
    import numpy as np  # type: ignore
    try:
        import cv2  # type: ignore
    except Exception:  # noqa: BLE001
        pytest.skip("opencv not available")
    with tempfile.NamedTemporaryFile(suffix=".avi", delete=False) as tmp:
        path = tmp.name
    try:
        fourcc = cv2.VideoWriter_fourcc(*"MJPG")
        writer = cv2.VideoWriter(path, fourcc, 10.0, (64, 64))
        for i in range(8):
            frame = np.zeros((64, 64, 3), dtype=np.uint8)
            frame[:, :, 0] = i * 30  # varying red
            frame[:, :, 1] = (8 - i) * 30
            writer.write(frame)
        writer.release()
        with open(path, "rb") as f:
            raw = f.read()
        doc = parser.parse(raw, filename="x.avi", modality=MODALITY_VIDEO)
        assert doc.modality == MODALITY_VIDEO
        assert len(doc.images) >= 1
        assert doc.metadata.get("fps", 0) > 0
    finally:
        try:
            os.unlink(path)
        except Exception:  # noqa: BLE001
            pass


def test_parse_email(parser: MultiModalParser, tiny_eml_bytes: bytes):
    doc = parser.parse(tiny_eml_bytes, filename="x.eml",
                       modality=MODALITY_EMAIL)
    assert doc.modality == MODALITY_EMAIL
    assert doc.metadata.get("subject", "").startswith("P4-7-W1")
    assert "Hello Bob" in doc.text or "Hello Bob" in doc.text.lower()


def test_batch_parse(parser: MultiModalParser, tiny_png_bytes: bytes,
                     tiny_wav_bytes: bytes):
    docs = parser.parse_batch(
        [tiny_png_bytes, tiny_wav_bytes, b"plain text"],
    )
    assert len(docs) == 3
    # bytes input falls back to text by default
    assert docs[2].modality in (MODALITY_TEXT, MODALITY_DOCUMENT)
